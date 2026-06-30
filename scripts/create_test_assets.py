import os
from docx import Document

def create_assets():
    target_dir = r"d:\MVNG\Lasya sree\Projects\AI in Recruitment\hiregenius AI\data\test_assets"
    os.makedirs(target_dir, exist_ok=True)

    # 1. Create Mock Resume DOCX
    resume_path = os.path.join(target_dir, "John_Doe_Resume.docx")
    doc_resume = Document()
    doc_resume.add_heading("John Doe", level=0)
    doc_resume.add_paragraph("Email: johndoe@gmail.com | Phone: +1-555-0199 | Location: New York, NY")
    
    doc_resume.add_heading("Professional Summary", level=1)
    doc_resume.add_paragraph(
        "Highly experienced Senior Full Stack Engineer with 5 years of software development experience. "
        "Specializes in Python backend architectures, FastAPI, relational databases, and React frontends."
    )
    
    doc_resume.add_heading("Technical Skills", level=1)
    p_skills = doc_resume.add_paragraph()
    p_skills.add_run("Backend: ").bold = True
    p_skills.add_run("Python, FastAPI, Django, Flask, SQL, MySQL, PostgreSQL\n")
    p_skills.add_run("Frontend: ").bold = True
    p_skills.add_run("JavaScript, TypeScript, React, HTML5, CSS3, Tailwind CSS\n")
    p_skills.add_run("DevOps & Cloud: ").bold = True
    p_skills.add_run("AWS, Docker, Git, CI/CD pipelines")

    doc_resume.add_heading("Professional Experience", level=1)
    doc_resume.add_heading("Senior Software Engineer | TechCorp (2023 - Present)", level=2)
    doc_resume.add_paragraph(
        "Led a team of developers in building and deploying high-performance microservices. "
        "Designed and implemented RESTful APIs using FastAPI and managed MySQL query optimization."
    )
    
    doc_resume.add_heading("Software Developer | DevStudio (2021 - 2023)", level=2)
    doc_resume.add_paragraph(
        "Developed web applications with React.js. Created robust backend features in Python and "
        "deployed containers using Docker on AWS."
    )
    
    doc_resume.save(resume_path)
    print(f"Created mock resume: {resume_path}")

    # 2. Create Mock Policy DOCX
    policy_path = os.path.join(target_dir, "Remote_Work_Policy.docx")
    doc_policy = Document()
    doc_policy.add_heading("HireGenius AI — Remote Work Policy", level=0)
    
    doc_policy.add_heading("1. Policy Overview", level=1)
    doc_policy.add_paragraph(
        "This policy outlines the guidelines and requirements for remote work arrangements at HireGenius AI. "
        "The company supports remote working arrangements to enhance flexibility and productivity."
    )

    doc_policy.add_heading("2. Eligibility & Hours", level=1)
    doc_policy.add_paragraph(
        "Full-time employees who have completed their probation period are eligible for remote work. "
        "Employees can request to work remotely for up to 3 days per week. "
        "Core working hours are defined as 10:00 AM to 4:00 PM EST, during which all remote employees must "
        "be reachable via email, Slack, and Zoom."
    )

    doc_policy.add_heading("3. Home Office Requirements", level=1)
    doc_policy.add_paragraph(
        "Remote employees are responsible for ensuring they have a secure, quiet workspace and a stable, "
        "high-speed internet connection (minimum 50 Mbps download speed). All corporate assets used remotely "
        "must be operated in compliance with our IT Security Guidelines."
    )

    doc_policy.save(policy_path)
    print(f"Created mock policy: {policy_path}")

if __name__ == "__main__":
    create_assets()
